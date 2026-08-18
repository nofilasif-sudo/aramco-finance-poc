"""Tests for the Word reader (word.py).

Plain unittest, no pytest needed:  python -m unittest discover -s tests

The bold tests carry the weight here. Bold is the only signal separating a
subtotal from a component in these documents, and every way of getting it
wrong fails SILENTLY — you get a full table of plausible rows with the
totals mislabelled.
"""

from __future__ import annotations

import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from fs_ingest.errors import IngestError                      # noqa: E402
from fs_ingest.word import read_tables                        # noqa: E402


def put(cell, text: str, bold: bool | None = None) -> None:
    """Write one run into a cell. bold=None leaves it inheriting."""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    if bold is not None:
        run.bold = bold


def save(document) -> Path:
    path = Path(tempfile.mkdtemp()) / "fixture.docx"
    document.save(path)
    return path


class TestBold(unittest.TestCase):
    def test_explicit_false_is_not_bold(self):
        """THE TRAP. python-docx writes bold=False as <w:b w:val="0"/>, so a
        reader testing for the tag's PRESENCE marks this row bold. If this
        test ever passes trivially, check that the fixture really is
        emitting val="0" and not omitting the tag."""
        doc = Document()
        doc.add_paragraph("Statement")
        table = doc.add_table(rows=3, cols=2)
        put(table.cell(0, 0), "")
        put(table.cell(0, 1), "Q1 2026")
        put(table.cell(1, 0), "Purchases", bold=False)
        put(table.cell(1, 1), "(111,692)", bold=False)
        put(table.cell(2, 0), "Operating costs", bold=True)
        put(table.cell(2, 1), "(244,700)", bold=True)

        path = save(doc)
        # A .docx is a ZIP, so the markup has to be read out of the entry
        # rather than off the file. Asserting on it keeps this test honest:
        # if python-docx ever switched to OMITTING the tag for non-bold, the
        # trap would no longer be exercised and the test below would start
        # passing for the wrong reason.
        with zipfile.ZipFile(path) as archive:
            markup = archive.read("word/document.xml").decode("utf-8")
        self.assertIn('w:val="0"', markup,
                      "fixture must exercise the explicit-false encoding")

        rows = read_tables(path)[0].rows
        self.assertFalse(rows[0].bold, "explicit val=0 must not read as bold")
        self.assertTrue(rows[1].bold)

    def test_mixed_runs_are_not_bold(self):
        """`all`, not `any`: a partly-bold caption is not a subtotal, and
        treating it as one drops a real line item out of its components."""
        doc = Document()
        doc.add_paragraph("Statement")
        table = doc.add_table(rows=2, cols=2)
        put(table.cell(0, 0), "")
        put(table.cell(0, 1), "Q1 2026")
        cell = table.cell(1, 0)
        put(cell, "Revenue ", bold=True)
        put(cell, "(Note 10)", bold=False)
        put(table.cell(1, 1), "433,101", bold=False)

        self.assertFalse(read_tables(save(doc))[0].rows[0].bold)

    def test_bold_inherited_from_paragraph_style(self):
        """run.bold is None means "inherit". python-docx does not walk the
        chain, and bool(None) is False — so an export that used a named
        "Total" style instead of direct formatting would quietly turn every
        subtotal into a component row."""
        doc = Document()
        style = doc.styles.add_style("TotalRow", WD_STYLE_TYPE.PARAGRAPH)
        style.font.bold = True
        doc.add_paragraph("Statement")
        table = doc.add_table(rows=2, cols=2)
        put(table.cell(0, 0), "")
        put(table.cell(0, 1), "Q1 2026")
        cell = table.cell(1, 0)
        cell.paragraphs[0].style = style
        put(cell, "Net income")            # bold left as None
        put(table.cell(1, 1), "122,008")

        self.assertTrue(read_tables(save(doc))[0].rows[0].bold)

    def test_empty_cell_is_not_bold(self):
        doc = Document()
        doc.add_paragraph("Statement")
        table = doc.add_table(rows=2, cols=2)
        put(table.cell(0, 0), "")
        put(table.cell(0, 1), "Q1 2026")
        put(table.cell(1, 1), "1")
        self.assertFalse(read_tables(save(doc))[0].rows[0].bold)


class TestText(unittest.TestCase):
    def test_fragmented_runs_are_joined(self):
        """Word splits a visible string across runs when spellcheck or a
        revision id touches it. Reading runs[0] would truncate the label."""
        doc = Document()
        doc.add_paragraph("Statement")
        table = doc.add_table(rows=2, cols=2)
        put(table.cell(0, 0), "")
        put(table.cell(0, 1), "Q1 2026")
        cell = table.cell(1, 0)
        for fragment in ("Borrow", "ings ", "(Note 9)"):
            put(cell, fragment, bold=False)
        put(table.cell(1, 1), "56,166", bold=False)

        self.assertEqual(read_tables(save(doc))[0].rows[0].cells[0],
                         "Borrowings (Note 9)")

    def test_horizontal_merge_collapses_to_one_cell(self):
        """python-docx yields the same cell once per grid column it spans."""
        doc = Document()
        doc.add_paragraph("Statement")
        table = doc.add_table(rows=2, cols=3)
        put(table.cell(0, 0), "")
        put(table.cell(0, 1), "Non-current")
        put(table.cell(0, 2), "Current")
        merged = table.cell(1, 1).merge(table.cell(1, 2))
        put(table.cell(1, 0), "Lease liabilities", bold=False)
        put(merged, "52,083", bold=False)

        row = read_tables(save(doc))[0].rows[0]
        self.assertEqual(len(row.cells), 2,
                         f"merge not collapsed: {row.cells}")


class TestStructure(unittest.TestCase):
    def test_heading_is_attached_and_not_reused(self):
        doc = Document()
        doc.add_paragraph("Consolidated balance sheet")
        first = doc.add_table(rows=2, cols=2)
        put(first.cell(0, 0), "")
        put(first.cell(0, 1), "31 Mar 2026")
        put(first.cell(1, 0), "Inventories", bold=False)
        put(first.cell(1, 1), "93,814", bold=False)
        doc.add_paragraph("Note 9 — Borrowings")
        second = doc.add_table(rows=2, cols=2)
        put(second.cell(0, 0), "")
        put(second.cell(0, 1), "Total")
        put(second.cell(1, 0), "Lease liabilities", bold=False)
        put(second.cell(1, 1), "64,062", bold=False)

        tables = read_tables(save(doc))
        self.assertEqual([t.title for t in tables],
                         ["Consolidated balance sheet", "Note 9 — Borrowings"])

    def test_table_without_heading_raises(self):
        """A silently mis-attributed statement would corrupt every
        downstream check, so this fails rather than guessing."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        put(table.cell(0, 0), "")
        put(table.cell(0, 1), "Q1 2026")
        put(table.cell(1, 0), "Orphan", bold=False)
        put(table.cell(1, 1), "1", bold=False)

        with self.assertRaises(IngestError):
            read_tables(save(doc))

    def test_second_table_does_not_inherit_first_heading(self):
        doc = Document()
        doc.add_paragraph("Only heading")
        for _ in range(2):
            table = doc.add_table(rows=2, cols=2)
            put(table.cell(0, 0), "")
            put(table.cell(0, 1), "Q1 2026")
            put(table.cell(1, 0), "Line", bold=False)
            put(table.cell(1, 1), "1", bold=False)

        with self.assertRaises(IngestError):
            read_tables(save(doc))

    def test_document_with_no_tables_raises(self):
        doc = Document()
        doc.add_paragraph("Nothing but prose")
        with self.assertRaises(IngestError):
            read_tables(save(doc))


if __name__ == "__main__":
    unittest.main()
