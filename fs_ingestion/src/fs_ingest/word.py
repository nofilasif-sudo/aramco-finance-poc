"""Reading tables out of a Word document a human authored.

The `.docx` counterpart to bronze_ingest's `excel.py`, and scoped the same
way: nothing here knows about financial statements or bronze. It solves one
problem — a Word file interleaves headings and tables, and the formatting a
reader's eye relies on (bold) is buried in run properties — and hands back
plain data.

Why not docx2txt / Pandoc / a PDF round-trip: all of them discard bold, and
bold is the only signal separating a subtotal row from a component row in
these documents. Most of the totals in the Group FS pack do not contain the
word "Total" ("Operating income", "Net income", "Net book value", "External
revenue"), so an unguarded SUM over the components double-counts and nothing
in the text would have warned you.

THE TRAP, and the reason this is a module rather than a dozen inline lines:
Word writes non-bold runs as an explicit `<w:b w:val="0"/>`, not by omitting
the tag. The Group FS pack holds 168 of those against 76 real `val="1"`.
Testing for the tag's PRESENCE therefore marks every row bold, which
silently destroys subtotal detection while appearing to work. python-docx
models this correctly as a tri-state (True / False / None), which is why we
lean on it instead of parsing the XML by hand.
"""

from __future__ import annotations

from dataclasses import dataclass

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as _DocxTable
from docx.text.paragraph import Paragraph as _DocxParagraph

from .errors import IngestError

# A style chain terminates in a handful of hops; anything longer is a cycle
# in the document's style table and we would rather stop than hang.
_MAX_STYLE_DEPTH = 16


@dataclass(frozen=True)
class WordRow:
    """One table row. `cells` is verbatim text; `bold` describes cells[0]."""
    cells: tuple[str, ...]
    bold: bool


@dataclass(frozen=True)
class WordTable:
    """One table, plus the heading paragraph that introduced it."""
    title: str
    header: tuple[str, ...]
    rows: tuple[WordRow, ...]


# ---------------------------------------------------------------------------
# Cell primitives
# ---------------------------------------------------------------------------
def cell_text(cell) -> str:
    """All text in a cell, joined and stripped.

    Runs are joined rather than read individually because Word fragments a
    single visible string across several runs whenever spellcheck or a
    revision id touches it — "Borrowings (Note 9)" can arrive as three runs.
    Reading runs[0] would silently truncate the label, and a truncated label
    is exactly the class of defect this pipeline exists to surface.
    """
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _style_bold(style) -> bool | None:
    """Walk a style and its ancestors for an explicit bold setting."""
    depth = 0
    while style is not None and depth < _MAX_STYLE_DEPTH:
        if style.font.bold is not None:
            return style.font.bold
        style = style.base_style
        depth += 1
    return None


def run_is_bold(run, paragraph) -> bool:
    """Resolve one run's bold to a definite True/False.

    `run.bold` is tri-state: True and False are direct formatting, None means
    "inherit". python-docx does NOT walk the inheritance chain, so None has to
    be resolved here — character style first, then paragraph style, then
    Word's document defaults (not bold).

    The Group FS pack sets bold directly on every run, so the None branch is
    unexercised by those two files. It is here because a pack exported with a
    named "Total" paragraph style rather than direct formatting would resolve
    every row to None, and `bool(None)` is False — every subtotal would
    quietly become a component row.
    """
    if run.bold is not None:
        return run.bold
    for style in (run.style, paragraph.style):
        resolved = _style_bold(style)
        if resolved is not None:
            return resolved
    return False


def cell_is_bold(cell) -> bool:
    """True when every text-bearing run in the cell is bold.

    `all`, not `any`: in a financial statement a partly-bold caption is not a
    subtotal, and treating it as one would drop a real line item out of its
    section's components. A cell with no text is not bold.
    """
    runs = [(r, p) for p in cell.paragraphs for r in p.runs if r.text.strip()]
    return bool(runs) and all(run_is_bold(r, p) for r, p in runs)


def row_cells(row) -> list:
    """A row's cells with horizontal merges collapsed to one entry.

    python-docx yields the SAME underlying cell once per grid column it
    spans, so a merged cell would otherwise read as several identical columns
    and inflate the row count. Identity of the `<w:tc>` element is the
    reliable discriminator; cell text is not, because two genuinely distinct
    cells can hold the same string.
    """
    seen, out = set(), []
    for cell in row.cells:
        key = id(cell._tc)
        if key not in seen:
            seen.add(key)
            out.append(cell)
    return out


# ---------------------------------------------------------------------------
# Document walk
# ---------------------------------------------------------------------------
def _blocks(document):
    """Paragraphs and tables in document order.

    `document.paragraphs` and `document.tables` are separate sequences with
    no way to interleave them, so a table cannot be matched to the heading
    above it through the public API. Walking the body's XML children is the
    documented way round it.
    """
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield _DocxParagraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield _DocxTable(child, document)


def read_tables(path) -> list[WordTable]:
    """Every table in the document, each tagged with its heading.

    The heading is the nearest preceding non-empty paragraph, verbatim. A
    table with no heading above it raises rather than landing under an empty
    or inherited title: the heading is this pack's grouping key, and a
    silently mis-attributed statement would corrupt every downstream check.
    """
    document = Document(str(path))
    name = getattr(path, "name", str(path))
    tables: list[WordTable] = []
    title = ""

    for block in _blocks(document):
        if isinstance(block, _DocxParagraph):
            text = block.text.strip()
            if text:
                title = text
            continue

        if not title:
            raise IngestError(
                f"{name}: table {len(tables) + 1} has no heading paragraph "
                f"above it — cannot attribute its rows to a statement.")

        rows = list(block.rows)
        if not rows:
            raise IngestError(f"{name}: '{title}' has an empty table.")

        header = tuple(cell_text(c) for c in row_cells(rows[0]))
        body = tuple(
            WordRow(cells=tuple(cell_text(c) for c in row_cells(r)),
                    bold=cell_is_bold(row_cells(r)[0]))
            for r in rows[1:]
        )
        tables.append(WordTable(title=title, header=header, rows=body))
        # Consumed: the next table must find its own heading rather than
        # inheriting this one.
        title = ""

    if not tables:
        raise IngestError(f"{name}: no tables found.")
    return tables
